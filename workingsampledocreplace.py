
#finds the word sea, and replaces it with 'new text containing ocean'

from docx import Document
from docx import Document
from docx.shared import Inches
import json

import datetime
from copy import deepcopy

#https://stackoverflow.com/questions/24872527/combine-word-document-using-python-docx



import docx

import docx2txt




"""def getText(test1):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)"""

#doesn't work


print ('done')

"""
#doesn't work well
files = ['test1.docx', 'test2.docx']

def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save('test4.docx')

combine_word_documents(files)
print ('done')"""
