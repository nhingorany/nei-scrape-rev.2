
#finds the word sea, and replaces it with 'new text containing ocean'
#fuck that - reads document
#search for x replace with xyz - but not quite there yet. learn regular expressions

#NOT QUITE WORKING KEEP PLAYING: https://stackoverflow.com/questions/9222106/how-to-extract-information-between-two-unique-words-in-a-large-text-file

from docx import Document

from docx.shared import Inches
import json

import datetime
from copy import deepcopy

from docx import Document
import re


f = open('test.txt','r') 
print (f.text)



#document = Document('test.txt')

#print (document.paragraphs)

"""for p in document.paragraphs:
    print (p.text)"""

import docx

"""def getText(f):
    doc = Document(f)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

for para in document.paragraphs:
   print (para.text)"""

print ('hoho')


#book1, page1 = re.split("[-/]", para.text)


zz = re.findall('1.4(.|\n)*?)1.4', f.text, re.M)
print (zz)


match = re.search('ABC(.*)XYZ', f.text)
print (match)
print ('search')

again = re.findall(r'(?<=ABC).*?(?=XYZ)', f.text, re.M)
print (again)
print ('findal DOTALl')

real = re.findall('ABC(.*?)XYZ', f.text, re.S)
print (real)
print ('multi')


x = re.findall(r'[ABC](.*?)[XYZ]',f.text,re.MULTILINE)
print (x)
print ('muLTILINE')


real = re.findall('(1.4)(.*?)(1.4)', f.text, re.S)
print (real)
print ('onepointfour')


#.join(x).replace('\n',' ')



"""pattern = r"<person>(.*?)</person>"
re.findall(pattern, str, flags=0)

print (pattern)"""


#https://stackoverflow.com/questions/25353652/regular-expressions-extract-text-between-two-markers
#m = re.search(ur'person>(?P<text>.*?)location>', para.text)



"""document = Document('Test.docx')

#Dictionary = {‘sea’: “ocean”}

sections = document.sections
for section in sections:
    print(section.start_type)


for paragraph in document.paragraphs:
    if 'sea' in paragraph.text:
        print (paragraph.text)
        paragraph.text = 'new text containing ocean'



document.save('Test2.docx')
print  ('done')"""



"""from docx import Document

def matches_my_condition(line):
     Returns true or false if the given line should be added to the document
    # Which will return true if the word cake appears in the line
    # return 'cake' in line

line = input("enter section")

# Prepare document
document = Document()

with open('test.docx, 'r'):
    for line in test.readlines():
        if matches_my_condition(line):
            document.add_paragraph(line)

document.save('test2.docx')"""

#1.4 in : test2
#1.4 - limited content survey in test


files = ['test1.docx', 'test2.docx']

"""def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save('test4.docx')

combine_word_documents(files)
print ('done')"""
