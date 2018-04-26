

import time
from time import sleep
import datetime
from PIL import Image
import pytesseract
import argparse
import re
import os
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import NoSuchElementException
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.keys import Keys


#text input

print ("program searches vision apprisal in Portsmouth, Middletown, Tiverton, Newport, Narragansett; gets gis, RIDEM soils, wetlands and aerial")
print  ("if using Tiverton use full name e.g. road, way, lane")
print ("if using Portsmouth Middletown or Newport use abbrev. eg. RD, WAY, LN")
print ("litteral searches for address only - not for use with vacant land")


print ("program search land evidence in Portsmouth, Middletown, Tiverton, Narragansett and Newport; gets lot info and screenshot of up to 5 deeds")
address = input("Enter Address")
town = input("Enter Town, no spaces, and capitalize first letter")
state = input("Enter Two Letter State")

print("open town page" + town)

driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')

if town == 'Portsmouth':
    web = 'https://i2e.uslandrecords.com/RI/Portsmouth/D/Default.aspx'
elif town == 'Middletown':
    web = 'https://i2f.uslandrecords.com/RI/Middletown/D/Default.aspx'
elif town == 'Tiverton':
    web = 'https://i2f.uslandrecords.com/RI/Tiverton/D/Default.aspx'
elif town == 'Narragansett':
    web = 'https://i2e.uslandrecords.com/RI/Narragansett/D/Default.aspx'
elif town == 'Newport':
    web = 'https://i2f.uslandrecords.com/RI/Newport'
else:
    print ('you spelled the Town name wrong or did not capitalize dummy')

print ('searching land records in')
print (town)
print (web)

    



#START HERE FOR RIDEM GIS MAP ---------------------------------------------------------
#RIDEM WETLANDS - SOILS - FLOOD ZONE MAP

driver.get('http://ridemgis.maps.arcgis.com/apps/webappviewer/index.html?id=87e104c8adb449eb9f905e5f18020de5') #get webpage mapss
driver.maximize_window()
time.sleep(2)

wait = WebDriverWait(driver, 10)
wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="esri_dijit_Search_0_input"]')))  #waits address button

element = driver.find_element_by_xpath('//*[@id="esri_dijit_Search_0_input"]') #finds address
element.send_keys(address + " " + town + " " + state) #sends address
element.send_keys(Keys.RETURN) #hits enter


wait = WebDriverWait(driver, 10)
wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="jimu_dijit_CheckBox_68"]/div[1]')))  #waits for soils button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_68"]/div[1]') #find soils button
element.click() #click soils button


element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_130"]/div[1]') #find  2014 aerial
element.click() #click button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_55"]/div[1]') #find  RIEMA FLOOD ZONE
element.click() #click button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_87"]/div[1]') #find  WETLANDS
element.click() #click button


element = driver.find_element_by_xpath('/html/body/div[2]/div/div[2]/div[2]/div[11]/img') #find legend
element.click() #click button

time.sleep(2)


driver.get_screenshot_as_file('ridemgis' + '.png')   #saves first screnshot

print ('done saved screenshot RIDEM')
driver.quit()

#END RIDEM WETLANDS----------------------------------

#START HERE FOR FEMA MAP--------------------------------------------

from selenium import webdriver
driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.keys import Keys

driver.get('https://msc.fema.gov/portal/search')

driver.maximize_window()

elem = driver.find_element_by_id('txtfloodmapsearch')
elem.clear()
elem.send_keys(address + " " + town + " " + state)
elem.send_keys(Keys.RETURN)

driver.execute_script("window.scrollTo(0,2000);")

time.sleep(3)

driver.get_screenshot_as_file('fema_map.png')

driver.quit()

print ("Done FEMA")

#DONE FEMA-------------------------------------------------------

#START GIS-------------------------------------------------------



print("open town page" + town)

driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

if town == 'Portsmouth':
    gis = 'http://www.mainstreetmaps.com/ri/portsmouth/public.asp'
elif town == 'Tiverton':
    gis = 'http://www.mainstreetmaps.com/ri/tiverton/public.asp'
elif town == 'Newport':  #arcgis
    gis = 'http://newportri.maps.arcgis.com/apps/webappviewer/index.html?id=78f7978f5667474da76d2533481662e4'
elif town == 'Middletown':  #tighebond
    gis = 'http://hosting.tighebond.com/MiddletownRI_Public/index.html'
else:
    print ('you spelled the Town name wrong or did not capitalize dummy OR town not in gis index')

driver = webdriver.Firefox()
driver.maximize_window()

#town list
mainst = ["Portsmouth", "Tiverton", "Warren"]
arcgis = ["Newport"]
tighebond = ["Middletown"]

if town in mainst:
    print("Using MainStMaps GIS")
    driver.get(gis)
    print ("searching for")
    print (town)
    print (gis)
    element = driver.find_element_by_id('d_disc_ok')
    element.click()

    elem = driver.find_element_by_id('s_location')
    elem.clear()
    elem.send_keys(address)
    elem.send_keys(Keys.DOWN)
    elem.send_keys(Keys.RETURN)

    time.sleep(3)

    elem = driver.find_element_by_id('baselayers')
    elem.send_keys(Keys.DOWN)
    elem.send_keys(Keys.RETURN)

    time.sleep(3)

    driver.get_screenshot_as_file('tempgis.png')
    print ("tempgis screenshot saved")
    
elif town in arcgis:   #arcgis search
    print("Using ARCGIS GIS")
    driver.get(gis)
    print ("searching for")
    print (town)
    print (gis)
    time.sleep(3)
    element = driver.find_element_by_xpath('//*[@id="esri_dijit_Search_0_input"]')  #click search bar
    element.click()
    element.send_keys(address)
    #elem.send_keys(Keys.UP)
    element.send_keys(Keys.RETURN)
    time.sleep(2)

    element = driver.find_element_by_xpath('//*[@id="widgets_ZoomSlider_Widget_31"]/div[1]')  #one click zoom in
    element.click()
    element = driver.find_element_by_xpath('//*[@id="widgets_ZoomSlider_Widget_31"]/div[1]')  #one click zoom in
    element.click()
    element = driver.find_element_by_xpath('//*[@id="widgets_ZoomSlider_Widget_31"]/div[1]')  #one click zoom in
    element.click()
    time.sleep(3)
    
    driver.get_screenshot_as_file('tempgis.png')
    print ("tempgis screenshot saved")

elif town in tighebond:  #tighebond search   #lot more to to click easements, etc. LOT OF INFO
    print("Using tighebond GIS")
    driver.get(gis)
    print ("searching for")
    print (town)
    print (gis)
    time.sleep(3)
    element = driver.find_element_by_xpath('//*[@id="searchinput"]')  #click search bar
    element.click()
    element.send_keys(address)
    #elem.send_keys(Keys.UP)
    element.send_keys(Keys.RETURN)
    time.sleep(3)
    element = driver.find_element_by_xpath('//*[@id="tabbasemap"]/button/div')  #click layer bar
    element.click()
    element = driver.find_element_by_xpath('//*[@id="baseMapGallery"]/li[4]/a/img')  #click googlem map bar
    element.click()
    
    
    time.sleep(3)
    driver.get_screenshot_as_file('tempgis.png')
    print ("tempgis screenshot saved")

  
   
else:
    print("town not in gis list")

driver.quit()



#END GIS ---------------------------

#START VISION --------------------------------

driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

driver.get('http://www.vgsi.com/vision/applications/parceldata/RI/Home.aspx')
driver.maximize_window()

element = driver.find_element_by_partial_link_text(town)
element.click()

element = driver.find_element_by_id('MainContent_btnEnterOnlineDatabase')
element.click()

elem = driver.find_element_by_id('MainContent_txtSearchAddress')


elem.send_keys(address)
time.sleep(2)
elem.send_keys(Keys.UP)
time.sleep(2)
elem.send_keys(Keys.RETURN)


time.sleep(3)

###scroll and screenshots


driver.get_screenshot_as_file('1vision appraisal.png')

driver.execute_script("window.scrollTo(0,1000);")


driver.get_screenshot_as_file('2vision appraisal.png')

driver.execute_script("window.scrollTo(0,1200);")

time.sleep(3)

driver.execute_script("window.scrollTo(0,2400);")

driver.get_screenshot_as_file('3vision appraisal.png')

#get relevant info owner, co, zone, area, plat lot
#vars: own, co, ap, bp, zone, area


own = driver.find_element_by_id('MainContent_lblOwner').text
print (own)

co = driver.find_element_by_id('MainContent_lblCoOwner').text
print (co)

ap = driver.find_element_by_id('MainContent_lblMblu').text
print (ap)

bp1 = driver.find_element_by_id('MainContent_lblBp').text
print ("book and page_" + bp1)

zone = driver.find_element_by_id('MainContent_lblZone').text
print (zone)

ac = driver.find_element_by_id('MainContent_lblLndAcres').text
print (ac)

#END VISION -----------------------------------------------------------------------

#splits book page with / - as delimiters

#BOOKPAGE1   #if element greater than 0, get element and parse by delim. if not exist print no data

if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[2]/td[4]')) > 0: 
    bp1 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[2]/td[4]').text
    print (bp1)
else:
    bp1 = "x"
    book1 = "x"
    page1 = "x"
    print ('no bp1')

if len(bp1) >= 3:    
    book1, page1 = re.split("[-/]", bp1)
    print ("book_" + book1)
    print ("page_" + page1)
else:
    print ('no data')
    
#BOOKPAGE2
if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[3]/td[4]')) > 0: 
    bp2 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[3]/td[4]').text
    print (bp2)
else:
    bp2 = "x"
    book2 = "x"
    page2 = "x"
    print ('no bp2')

if len(bp2) >= 3:    
    book2, page2 = re.split("[-/]", bp2)
    print ("book_" + book2)
    print ("page_" + page2)
else:
    print ('no data')
    
#BOOKPAGE3
if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[4]/td[4]')) > 0: 
    bp3 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[4]/td[4]').text
    print (bp3)
else:
    bp3 = "x"
    book3 = "x"
    page3 = "x"
    print ('no bp3')

if len(bp3) >= 3:    
    book3, page3 = re.split("[-/]", bp3)
    print ("book_" + book3)
    print ("page_" + page3)
else:
    print ('no data')



#BOOKPAGE4
   
if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]')) > 0: 
    bp4 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]').text
    print (bp4)
else:
    bp4 = "x"
    book4 = "x"
    page4 = "x"
    print ('no bp4')

if len(bp4) >= 3:    
    book4, page4 = re.split("[-/]", bp4)
    print ("book_" + book4)
    print ("page_" + page4)
else:
    print ('no data')

#BOOKPAGE5
    
if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]')) > 0: 
    bp5 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]').text
    print (bp5)
else:
    bp5 = "x"
    book5 = "x"
    page5 = "x"
    print ('no bp4')
    

if len(bp5) >= 3:    
    book5, page5 = re.split("[-/]", bp4)
    print ("book_" + book5)
    print ("page_" + page5)
else:
    print ('no data')

driver.quit()

#END VISION ---------------------

#IMAGE CROPPING

#FEMA CROP

from PIL import Image
 
""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('fema_map.png')
box = (450, 0, 1500, 900)
area = img.crop(box)

area.save('cropfema.png')

#VISION CROP1

""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('1vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save('1VAcrop.png')

#VISION CROP2

""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('2vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save('2VAcrop.png')

#VISION CROP3

""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('3vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save('3VAcrop.png')



#END CROPPING

#START HERE FOR WORD DOC CREATION



#WORD DOCUMENT START ---------------------------------------------------------

from docx import Document
from docx.shared import Inches
import json

import datetime

now = datetime.datetime.now()

document = Document()

from docx.shared import Pt

#NEW 04.21.18 - INSERT SECTIONS FROM WORD DOC:

import docx
import re


#reads docx and converts to text variable fullText
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

if __name__ == '__main__':
    filename='NEI - MASTER PROPOSAL SCOPE - HOURLY RATES - NOTES 2018.docx';  #docx file name
    fullText=getText(filename)
    # print (fullText)  #USE IF YOU WANT TO PRINT ENTIRE DOCUMENT

#print (fullText)

#finds text between two markers, 101, 101a FROM MASTER WORD DOCUMENT
match1 = re.findall('101(.*?)101a', str(fullText), re.DOTALL)


#removes newline \n and prints TEXT!!!! FOR MARKED SECTIONS
p5 = re.sub("\n", " ", str(''.join(match1)), re.DOTALL)
print (p5)





#CREATE NEW DOCUMENT

style = document.styles['Normal']
font = style.font
font.name = 'Swis721 Lt BT'
font.size = Pt(10)

document.add_heading('Narragansett Engineering Inc.', 0)

from docx.enum.text import WD_LINE_SPACING

paragraph = document.add_paragraph()
paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

p = document.add_paragraph('3102 East Main Rd, Portsmouth RI 02871', style='No Spacing')
p = document.add_paragraph('T. 401.683.6630', style='No Spacing')
p = document.add_paragraph('nei-cds.com', style='No Spacing')

p = document.add_paragraph ('Date + Time: ' + now.strftime("%Y-%m-%d %H:%M"), style='No Spacing')


document.add_paragraph('to: ' + own + " " + co + " " + address + " " + town + " " + state, style='No Spacing')
document.add_paragraph('Plat + Lot (A.P.).: ' + ap, style='No Spacing')
document.add_paragraph('                                             ')
document.add_paragraph('Site Information: ' + bp1, style='No Spacing')
document.add_paragraph('Latest Book and Page: ' + bp1, style='No Spacing')
document.add_paragraph('Land evidence chain: ' + bp2 + " " + bp3 + " " + bp4 + " " + bp5, style='No Spacing')
document.add_paragraph('Zone: ' + zone, style='No Spacing')
document.add_paragraph('Lot Area (Acres): ' + ac + '+/-', style='No Spacing')
document.add_paragraph('All information from Assessors Database ', style='No Spacing')

document.add_paragraph('                                             ')
document.add_paragraph('                                             ')
document.add_paragraph('Narragansett Engineering Inc is please to provide you with the following proposal', style='No Spacing')




p = document.add_paragraph(town + ' GIS + Aerial Map: ')
r = p.add_run()
r.add_picture('tempgis.png', width=Inches(5.0))


p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_text('From assessors datatabase: ' + now.strftime("%Y-%m-%d %H:%M"))
r.add_picture('1VAcrop.png', width=Inches(4.0))

p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_picture('2VAcrop.png', width=Inches(4.0))

p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_picture('3VAcrop.png', width=Inches(4.0))


p = document.add_paragraph('RIDEM GIS Soils, Wetlands, Flood Zone: ')
r = p.add_run()
r.add_picture('ridemgis.png', width=Inches(4.0))

p = document.add_paragraph('FEMA GIS Information: ')
r = p.add_run()
r.add_picture('cropfema.png', width=Inches(4.0))


document.add_page_break()


p = document.add_paragraph(p5)  #adds new paragraph from MASTER DOCUMENT
r = p.add_run()


document.save('NEI SAMPLE.docx')

print ("Done! Check NEI Sample.docx")



#WORD DOCUMENT END ---------------------------------------------------------

#START COPY LOOP 1 -----------------   SEARCHES LAND EVIDENCE

print ('back to  main window')

if bp1 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book1) #book1                                             #XXX

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page xx
    elem.send_keys(page1)   #page1                                                          #XXX
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
          
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file(bp1 + '_deed.png')  #cdeed round 1 xx
        print ("screenshot_" + bp1)                         #  XXX

        time.sleep(5)

        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  bp1 + "_" + str(counter) + "_deed" + ".png"   #save as filename xx
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book1 + " " + page1) #XX
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp2")
    driver.quit()
    

#END COPY LOOP 1  ---------------


#START COPY LOOP 2 -----------------   

print ('back to  main window')

if bp2 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book2) #book2                                             #XXX

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page xx
    elem.send_keys(page2)   #page2                                                          #XXX
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
          
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file(bp2 + '_deed.png')  #cdeed round 2 xx
        print ("screenshot_" + bp2)                         #  XXX

        time.sleep(5)

        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  bp2 + "_" + str(counter) + "_deed" + ".png"   #save as filename xx
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book2 + " " + page2) #XX
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp2")
    driver.quit()
    

#END COPY LOOP 2  ---------------


#START COPY LOOP 3 -----------------   4 first to test window


print ('back to  main window')

if bp3 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(3)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book3) #book3

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page
    elem.send_keys(page3)   #page3
    elem.send_keys(Keys.RETURN)

    time.sleep(5)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(3)
            
        # Get windows list and put focus on new window (which is on the 1st index in the list)     
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file(bp3 + '_deed.png')  #cdeed round 3
        print ("screenshot_" + bp3) #save bp3

        time.sleep(5)


        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  "book3_" + "page3_" + str(counter) + "_deed" + ".png"   #save as filename xx
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book3 + " " + page3)  ##XX
        
    print  ("done item 3")
    driver.quit()
    
else:
    print ("no value for bp3")
    driver.quit()

#END COPY LOOP 3  ---------------


#START COPY LOOP 4 -----------------   4 first to test window

print ('back to  main window')


if bp4 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book4) #book4

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page
    elem.send_keys(page4)   #page4
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
        
        #ROUND 4  'c' -----
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file(bp4 + '0Cdeed.png')  #cdeed round 4
        print ("screenshot_" + bp4)

        time.sleep(5)


        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  "book4_" + "page4_" + str(counter) + "_deed" + ".png"   #save as filename
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book4 + " " + page4)
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp4")
    driver.quit()

#END COPY LOOP 4  ---------------

    

#START COPY LOOP 5 -----------------

if bp5 != "x":     

    print ('back to  main window')

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book5) #book5                                             #XXX

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page xx
    elem.send_keys(page5)   #page5                                                          #XXX
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
          
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file(bp5 + '_deed.png')  #cdeed round 5 xx
        print ("screenshot_" + bp5)                         #  XXX

        time.sleep(5)

        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  "book5_" + "page5_" + str(counter) + "_deed" + ".png"   #save as filename xx
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book5 + " " + page5) #XX
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp5")
    driver.quit()
    

#END COPY LOOP 5  ---------------

print ("DONE!! - wow, neal is pretty great")
driver.quit()












#try timeout exception...
"""from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait

try:
    element = WebDriverWait(driver, 10).until(lambda driver: driver.find_element_by_id('user_first_name'))
    # do smth with the found element
except TimeoutException:
    print "Element Not Found"
    driver.close()"""



"""
browser = webdriver.Firefox()
browser.get('https://www.google.com?q=python#q=python')
first_result = ui.WebDriverWait(browser, 15).until(lambda browser: browser.find_element_by_class_name('rc'))
first_link = first_result.find_element_by_tag_name('a')

# Save the window opener (current window)
main_window = browser.current_window_handle

# Open the link in a new window by sending key strokes on the element
first_link.send_keys(Keys.SHIFT + Keys.RETURN)

# Get windows list and put focus on new window (which is on the 1st index in the list)
windows = browser.window_handles
browser.switch_to.window(windows[1])

# do whatever you have to do on this page, we will just got to sleep for now
sleep(2)

# Close current window
browser.find_element_by_tag_name('body').send_keys(Keys.CONTROL + 'w')

# Put focus back on main window
browser.switch_to.window(main_window)"""

"""# Opens a new tab
self.driver.execute_script("window.open()")

# Switch to the newly opened tab
self.driver.switch_to.window(self.driver.window_handles[1])

# Navigate to new URL in new tab
self.driver.get("https://google.com")
# Run other commands in the new tab here
You're then able to close the original tab as follows

# Switch to original tab
self.driver.switch_to.window(self.driver.window_handles[0])

# Close original tab
self.driver.close()

# Switch back to newly opened tab, which is now in position 0
self.driver.switch_to.window(self.driver.window_handles[0])
Or close the newly opened tab

# Close current tab
self.driver.close()

# Switch back to original tab
self.driver.switch_to.window(self.driver.window_handles[0])"""





"""
#kill firefox and gecko
if (browser == "FIREFOX")) {
    try {
        Runtime.getRuntime().exec("taskkill /F /IM geckodriver.exe");
        Runtime.getRuntime().exec("taskkill /F /IM plugin-container.exe");
        Runtime.getRuntime().exec("taskkill /F /IM firefox.exe");
    } catch (IOException e) {
        e.printStackTrace();
    }
} else {
    driver.quit();
}
driver.quit()"""

