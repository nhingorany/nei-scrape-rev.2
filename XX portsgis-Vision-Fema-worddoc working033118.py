
"""so far. gets vision appraisal data by town. gets relavant info (owner, zone, etc.)
copy said info to word, along with screenshots of va; also gets fema map and screenshot)"""

"""GETS INFORMATION FROM VISION APPRIASAL - SEARCHES BY TOWN, SAVES THREE SCREENSHOTS AFTER SCROLLING DOWN"""

"""GET THIS TO WORK LATER: https://ripropinfo.com/"""

#NEEDS. FIREFOX 57ISH, SELENIUM, PYTHON36, PIL, PYTHON DOCX
#IMPORT PACKAGES

import time
from time import sleep


from selenium import webdriver
driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.keys import Keys

import time
from time import sleep
import datetime

from PIL import Image

import re
import os

from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC 

from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import NoSuchElementException

from selenium.webdriver.common.keys import Keys

driver = webdriver.Firefox()
driver.maximize_window()



#INPUT ADDRESS, TOWN, STATE

address = input("Enter Address")
town = input("Enter Town, no spaces, and capitalize first letter")
state = input("Enter Two Letter State")

print("open town page" + town)

#PORTSMOUTH RI GIS ONLY. EXAPAND TO MAINST MAPS AND REST FROM HERE
#https://ripropinfo.com/


tiverton: http://www.mainstreetmaps.com/ri/tiverton/public.asp

if town == "Portsmouth":
    driver.get('http://www.mainstreetmaps.com/ri/portsmouth/public.asp')

    element = driver.find_element_by_id('d_disc_ok')
    element.click()

    elem = driver.find_element_by_id('s_location')
    elem.clear()
    elem.send_keys(address)
    elem.send_keys(Keys.UP)
    elem.send_keys(Keys.RETURN)

    elem = driver.find_element_by_id('baselayers')
    elem.send_keys(Keys.DOWN)
    elem.send_keys(Keys.RETURN)

    time.sleep(3)

    driver.get_screenshot_as_file('tempgis.png')

#END GIS PORTSMOUTH


#START HERE FOR VISION

driver.get('http://www.vgsi.com/vision/applications/parceldata/RI/Home.aspx')

element = driver.find_element_by_partial_link_text(town)
element.click()

element = driver.find_element_by_id('MainContent_btnEnterOnlineDatabase')
element.click()

elem = driver.find_element_by_id('MainContent_txtSearchAddress')


elem.send_keys(address)
time.sleep(2)
elem.send_keys(Keys.UP)
time.sleep(2)
elem.send_keys(Keys.RETURN)


time.sleep(2)

#kind of working


driver.get_screenshot_as_file('1vision appraisal.png')

driver.execute_script("window.scrollTo(0,1000);")


driver.get_screenshot_as_file('2vision appraisal.png')

driver.execute_script("window.scrollTo(0,1200);")

time.sleep(2)

driver.execute_script("window.scrollTo(0,2400);")

driver.get_screenshot_as_file('3vision appraisal.png')


#get relevant info owner, co, zone, area, plat lot
#vars: own, co, ap, bp, zone, area


own = driver.find_element_by_id('MainContent_lblOwner').text
print (own)

co = driver.find_element_by_id('MainContent_lblCoOwner').text
print (co)


ap = driver.find_element_by_id('MainContent_lblMblu').text
print (ap)


bp = driver.find_element_by_id('MainContent_lblBp').text
print (bp)

zone = driver.find_element_by_id('MainContent_lblZone').text
print (zone)

ac = driver.find_element_by_id('MainContent_lblLndAcres').text
print (ac)

book, page = bp.split("-")

print (book)  #splits book
print (page)  #splits page


#START HERE FOR FEMA MAP--------------------------------------------

from selenium import webdriver
driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.keys import Keys

driver.get('https://msc.fema.gov/portal/search')

driver.maximize_window()

elem = driver.find_element_by_id('txtfloodmapsearch')
elem.clear()
elem.send_keys(address + " " + town + " " + state)
elem.send_keys(Keys.RETURN)

driver.execute_script("window.scrollTo(0,2000);")

time.sleep(3)

driver.get_screenshot_as_file('fema_map.png')

driver.quit()

print ("Done Scraping")

#DONE FEMA-------------------------------------------------------


#START HERE FOR RIDEM GIS MAP ---------------------------------------------------------
#RIDEM WETLANDS - SOILS - FLOOD ZONE MAP

driver.get('http://ridemgis.maps.arcgis.com/apps/webappviewer/index.html?id=87e104c8adb449eb9f905e5f18020de5') #get webpage mapss
driver.maximize_window()
time.sleep(2)

wait = WebDriverWait(driver, 10)
wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="esri_dijit_Search_0_input"]')))  #waits address button

element = driver.find_element_by_xpath('//*[@id="esri_dijit_Search_0_input"]') #finds address
element.send_keys(address + " " + town + " " + state) #sends address
element.send_keys(Keys.RETURN) #hits enter


wait = WebDriverWait(driver, 10)
wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="jimu_dijit_CheckBox_68"]/div[1]')))  #waits for soils button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_68"]/div[1]') #find soils button
element.click() #click soils button


element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_130"]/div[1]') #find  2014 aerial
element.click() #click button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_55"]/div[1]') #find  RIEMA FLOOD ZONE
element.click() #click button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_87"]/div[1]') #find  WETLANDS
element.click() #click button


element = driver.find_element_by_xpath('/html/body/div[2]/div/div[2]/div[2]/div[11]/img') #find legend
element.click() #click button

time.sleep(2)


driver.get_screenshot_as_file(str(address) + " " + str(town) + 'ridemgis' + '.png')   #saves first screnshot

print ('done saved screenshot RIDEM')

#END RIDEM WETLANDS----------------------------------


#IMAGE CROPPING

#FEMA CROP

from PIL import Image
 
""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('fema_map.png')
box = (450, 0, 1500, 900)
area = img.crop(box)

area.save('cropfema.png')

#VISION CROP1

""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('1vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save('1VAcrop.png')

#VISION CROP2

""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('2vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save('2VAcrop.png')

#VISION CROP3

""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('3vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save('3VAcrop.png')

#GIS CROP - edit crop

"""img = Image.open('tempgis.png')
box = (250, 200, 1440, 950)
area = img.crop(box)

area.save('giscrop.png')"""

#needs work on crop area

print ("Done Cropping")

#START HERE FOR WORD DOC CREATION


"""http://python-docx.readthedocs.io/en/latest/user/install.html
use to install docx from pip"""

#CREATES STD WORD FILE AND INSERTS VISION APPRAISAL IMAGES---------------------------------------------





#WORD DOCUMENT START ---------------------------------------------------------


from docx import Document
from docx.shared import Inches
import json

import datetime

now = datetime.datetime.now()

document = Document()
from docx.shared import Pt

style = document.styles['Normal']
font = style.font
font.name = 'Swis721 Lt BT'
font.size = Pt(10)

document.add_heading('Narragansett Engineering Inc.', stlye= Heading 2)

from docx.enum.text import WD_LINE_SPACING

paragraph = document.add_paragraph()
paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

p = document.add_paragraph('3102 East Main Rd, Portsmouth RI 02871', style='No Spacing')
p = document.add_paragraph('T. 401.683.6630', style='No Spacing')
p = document.add_paragraph('nei-cds.com', style='No Spacing')

p = document.add_paragraph ('Date + Time: ' + now.strftime("%Y-%m-%d %H:%M"), style='No Spacing')


document.add_paragraph('to: ' + own + " " + co + " " + address + " " + town + " " + state, style='No Spacing')
document.add_paragraph('Plat + Lot (A.P.).: ' + ap, style='No Spacing')
document.add_paragraph('Latest Book and Page: ' + bp1, style='No Spacing')
document.add_paragraph('Land evidence chain: ' + bp2 + " " + bp3 + " " + bp4 + " " + bp5, style='No Spacing')
document.add_paragraph('Zone: ' + zone, style='No Spacing')
document.add_paragraph('Lot Area (Acres): ' + ac + '+/-', style='No Spacing')
document.add_paragraph('All information from Assessors Database ', style='No Spacing')


p = document.add_paragraph(town + ' GIS + Aerial Map: ')
r = p.add_run()
r.add_picture('tempgis.png', width=Inches(5.0))


p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_text('From assessors datatabase: ' + now.strftime("%Y-%m-%d %H:%M"))
r.add_picture('1VAcrop.png', width=Inches(4.0))

p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_picture('2VAcrop.png', width=Inches(4.0))

p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_picture('3VAcrop.png', width=Inches(4.0))


p = document.add_paragraph('RIDEM GIS Soils, Wetlands, Flood Zone: ')
r = p.add_run()
r.add_picture('ridemgis.png'', width=Inches(4.0))

p = document.add_paragraph('FEMA GIS Information: ')
r = p.add_run()
r.add_picture('cropfema.png', width=Inches(4.0))


document.add_page_break()

document.save('NEI SAMPLE.docx')

print ("Done! Check NEI Sample.docx")



#WORD DOCUMENT END ---------------------------------------------------------
