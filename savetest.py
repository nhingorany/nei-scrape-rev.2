import docx
import re
import os
from docx.enum.text import WD_LINE_SPACING

from docx import Document
from docx.shared import Inches
import json

import datetime
from copy import deepcopy



from docx.shared import Pt
from docx.shared import Pt
#reads docx and converts to text variable fullText
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

if __name__ == '__main__':
    filename='NEI - MASTER PROPOSAL SCOPE - HOURLY RATES - NOTES 2018.docx';  #docx file name
    fullText=getText(filename)
    # print (fullText)


##############################  going to want to loop this. can only do two sections FTM
    #need to add notes to match section no's like elv certs'


#FIRST SECTION
sect1 = input("Enter Section1 - three digit code from master")  #first section
match1 = re.findall(sect1 +'(.*?)' + sect1, str(fullText), re.DOTALL)  #adds section no.
print (match1)

p1 = re.sub("\n", " ", str(''.join(match1)), re.DOTALL)  #removes newline section
print (p1)


#STANDARD NOTES
sect0 = "s.00"   #standard notes
match0 = re.findall(sect0 +'(.*?)' + sect0, str(fullText), re.DOTALL)  #adds standard notes.


    

p0 = re.sub("\n", " ", str(''.join(match0)), re.DOTALL)  #removes newline section
print (p0)

#SECTION NOTES

matchlist1 = ["101", "102", "103", "104", "105", "106", "107", "108"]  #MATCHLIST NO SECTION 1
matchlist2 = ["201", "202", "203", "204", "205" ,"206" , "207", "208"]  #MATCHLIST NO SECTION 2
matchlist3 = ["301", "302", "303" ,"304" ,"305" ,"306" ,"307"]

if sect1 in matchlist1: #if section 1, then section 1 std notes
    sectN1 = "s.012"
    print (sectN1)
elif sect1 in  matchlist2: #if section 1, then section 1 std notes
    sectN1 = "s.012"
    print (sectN1)
elif sect1 in  matchlist3: #if section 1, then section 1 std notes
    sectN1 = "s.03"
    print (sectN1)
else:
    print ('not in list USER ERROR')

    
match3 = re.findall(sectN1 +'(.*?)' + sectN1, str(fullText), re.DOTALL)  #adds section no. PER FIRST SECTION
print (match3)

p3 = re.sub("\n", " ", str(''.join(match3)), re.DOTALL)  #removes Newline /n for NOTES FOR SECTION
print (p3)


#additional section
cont = input ('enter Y/N')

if cont == "Y":
    sect2 = input ('enter sect2')
    match2 = re.findall(str(sect2) +'(.*?)' + str(sect2), str(fullText), re.DOTALL)  #adds section no.
    p2 = re.sub("\n", " ", str(''.join(match2)), re.DOTALL)
    print ('sect2')
    print (sect2)

if cont == "N":
    sect2 = "0"
    p2 = "0"
    print ('ok')
else:
    sect2 = "0"
    p2 = "0"
    print ('bad input - done')


    
    

###################


#print (fullText)

#finds text between two markers, [INPUT SECTION]
"""match1 = re.findall('101(.*?)101a', str(fullText), re.DOTALL)


#removes newline \n and prints TEXT!!!!
p5 = re.sub("\n", " ", str(''.join(match1)), re.DOTALL)
print (p5)


#-------------------------------------------------------------

#finds text between two markers, s.00s, s.00e
match2 = re.findall('s.00s(.*?)s.00e', str(fullText), re.DOTALL)


#removes newline \n and prints TEXT!!!!
p6 = re.sub("\n", " ", str(''.join(match2)), re.DOTALL)
print (p6)"""


#style = document.styles['Normal']
#font = style.font
#font.name = 'Swis721 Lt BT'
#font.size = Pt(10)

now = datetime.datetime.now()

document = Document()

paragraph = document.add_paragraph()

document.add_heading('Narragansett Engineering Inc.', 0)


document.add_page_break()

#####################reuse

p = document.add_paragraph(p1)  #adds new SECTION SCOPE from MASTER DOCUMENT
r = p.add_run()

if len(sect2) >= 3:
    p = document.add_paragraph(p2)  #adds standard notes
    r = p.add_run()
    print ('adding second section')
else:
    print ('nothing to add')

p = document.add_paragraph(p0)  #adds standard notes
r = p.add_run()

p = document.add_paragraph(p3)  #adds section notes
r = p.add_run()


##################reuse

direct = '104 Dighton'
addss =  'portsmouth'

import pathlib
pathlib.Path('c:/python36/docs/' + addss + '/').mkdir(parents=True, exist_ok=True)

document.save('c:/python36/docs/' + addss + '/' + direct + 'NEI SAMPLE2.docx')

print ("Done! Check NEI Sample.docx")
print ("saved in " + 'c:/python36/docs/' + addss + '/' + direct + 'NEI SAMPLE2.docx')
