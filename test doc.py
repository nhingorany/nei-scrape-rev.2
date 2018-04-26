from docx import Document
from docx.shared import Inches
import json

import datetime
from copy import deepcopy

now = datetime.datetime.now()

document = Document()

from docx.shared import Pt
from docx.shared import Pt


style = document.styles['Normal']
font = style.font
font.name = 'Swis721 Lt BT'
font.size = Pt(9)

own = "Neal"
co =  "Co owner"
address = "104 Dighton"
town = "Portsmouth"
state = "RI"
ap = "01-01"
zone = "R"
ac = "1.0"
bp1 = "1"
bp2 = "2"
bp3 = "3"
bp4 = "4"
bp5 = "5"



document.add_heading('Narragansett Engineering Inc.', 0)

from docx.enum.text import WD_LINE_SPACING

paragraph = document.add_paragraph()
paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

p = document.add_paragraph()
r = p.add_run()
r.add_picture('c:/python36/NEI Logo-Proposal.jpg', width=Inches(5.0))


p = document.add_paragraph('3102 East Main Rd, Portsmouth RI 02871') #, style='No Spacing')
p = document.add_paragraph('T. 401.683.6630')
p = document.add_paragraph('nei-cds.com')

p = document.add_paragraph ('Date + Time: ' + now.strftime("%Y-%m-%d %H:%M"))


document.add_paragraph('to: ' + own + " " + co + " " + address + " " + town + " " + state)
document.add_paragraph('Plat + Lot (A.P.).: ' + ap)
document.add_paragraph('                                             ')
document.add_paragraph('Site Information: ' + bp1)
document.add_paragraph('Latest Book and Page: ' + bp1)
document.add_paragraph('Land evidence chain: ' + bp2 + " " + bp3 + " " + bp4 + " " + bp5)
document.add_paragraph('Zone: ' + zone)
document.add_paragraph('Lot Area (Acres): ' + ac + '+/-')
document.add_paragraph('All information from Assessors Database ')

document.add_paragraph('                                             ')
document.add_paragraph('                                             ')
document.add_paragraph('Narragansett Engineering Inc is please to provide you with the following proposal')


print ("done")


