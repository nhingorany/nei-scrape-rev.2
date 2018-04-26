from docx import Document
import re
import docx
import os
import word, pdf

document = Document('NEI - MASTER PROPOSAL SCOPE - HOURLY RATES - NOTES 2018.docx')


"""def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

for para in document.paragraphs:
   print (para.text)"""


def getTextWord(wordFileName):
    doc = docx.Document(wordFileName)
    fulltext = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
            
    
def getText(txtFileName):
    file = open(textFileName, 'r')
    return file.read()



"""for para in document.paragraphs:
    match1 = re.findall('101(.*?)101a', para.text, re.MULTILINE)
    print (match1)"""

"""getText(filename)()
d = para.text
print (d)"""

"""doc = docx.Document(filename)
for p in doc.paragraphs:
print (x)"""



"""z = document.paragraphs
print (z.text)

def f():
    for p in document.paragraphs:
        global x
        x = p.text    
        print (x)

print ('next step')"""




        
    
        
        
    




"""
master()
match1 = re.findall('101(.*?)101a', y, re.MULTILINE)
print (match1)"""





