import re

import docx

from docx import Document

document = Document('NEI - MASTER PROPOSAL SCOPE - HOURLY RATES - NOTES 2018.docx')



def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

for para in document.paragraphs:
   print (para.text)
   match2 = re.findall('101($.*?)101a', para.text, flags=re.S)
   print (match2)
   


print('GAAAAAAAAAHHHHHHHHHHHHH')

   
print(para.text)
